import os
import datetime
import tempfile

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

from googleapiclient.http import MediaFileUpload

from utils.html_utils import html_to_lines, sanitize_filename

import os
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def disable_table_autofit(table):
    """Memaksa tabel menggunakan layout tetap (Fixed Layout) dan mematikan AutoFit."""
    tblPr = table._tbl.tblPr
    # Cari atau buat elemen w:tblLayout
    tblLayout = tblPr.find(qn('w:tblLayout'))
    if tblLayout is None:
        tblLayout = OxmlElement('w:tblLayout')
        tblPr.append(tblLayout)
    # Set type menjadi 'fixed' agar kolom patuh pada Cm() yang ditentukan
    tblLayout.set(qn('w:type'), 'fixed')

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Mengatur padding internal (margins) pada sel tabel."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for margin, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{margin}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_surat_pengantar(form_data, output_path="surat_pengantar.docx"):
    doc = Document()

    # 1. Pengaturan Margin Halaman (1 Inci di semua sisi)
    for section in doc.sections:
        section.top_margin = Cm(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

# 2. KOP SURAT (Menggunakan Tabel tanpa border untuk presisi Logo & Teks)
    kop_table = doc.add_table(rows=1, cols=2)
    kop_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Mematikan AutoFit agar Word patuh pada ukuran kolom statis
    disable_table_autofit(kop_table)
    
    # Pengaturan Lebar Kolom: Kolom 1 (20% = 3.3 Cm), Kolom 2 (80% = 13.2 Cm)
    kop_table.columns[0].width = Cm(3.3)
    kop_table.columns[1].width = Cm(13.2)
    kop_table.rows[0].cells[0].width = Cm(3.3)
    kop_table.rows[0].cells[1].width = Cm(13.2)

    # Masukkan Logo Kejaksaan
    cell_logo = kop_table.cell(0, 0)
    p_logo = cell_logo.paragraphs[0]
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Cek apakah file logo ada, jika tidak gunakan teks placeholder
    if os.path.exists("static/images/logo-kejaksaan.png"):
        p_logo.add_run().add_picture("static/images/logo-kejaksaan.png", width=Cm(2.5))
    else:
        p_logo.add_run("[ LOGO ]").bold = True

    # Masukkan Teks Kop Kejaksaan
    cell_text = kop_table.cell(0, 1)
    p_text = cell_text.paragraphs[0]
    p_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    r1 = p_text.add_run("KEJAKSAAN REPUBLIK INDONESIA\n")
    r1.font.size = Pt(14)
    r1.bold = True
    
    r2 = p_text.add_run("KEJAKSAAN TINGGI BALI\n")
    r2.font.size = Pt(16)
    r2.bold = True
    
    r3 = p_text.add_run("KEJAKSAAN NEGERI TABANAN\n")
    r3.font.size = Pt(24)
    r3.bold = True
    
    r4 = p_text.add_run(
        "Jalan Sudirman No. 5 Kabupaten Tabanan 82114\n"
        "Telp. (0361) 811083 Fax. (0361) 811325 www.kejari-tabanan.kejaksaan.go.id"
    )
    r4.font.size = Pt(9)

    # Garis Pembatas Kop Surat
    p_line = doc.add_paragraph()
    p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_line.paragraph_format.space_after = Pt(0)
    if os.path.exists("static/images/line.png"):
        p_line.add_run().add_picture("static/images/line.png", width=Cm(16.5))
    else:
        # Fallback jika gambar line.png belum ditaruh di folder
        p_line.add_run("=======================================================").bold = True

    # 3. TANGGAL SURAT (Rata Kanan)
    p_tgl = doc.add_paragraph()
    p_tgl.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_tgl.paragraph_format.space_after = Pt(12)
    p_tgl.add_run(f"Singasana, {form_data['tanggal_surat']}")

    # 4. TUJUAN SURAT (Kepada Yth.)
    p_tujuan = doc.add_paragraph()
    p_tujuan.paragraph_format.line_spacing = 1.15
    p_tujuan.paragraph_format.space_after = Pt(18)
    p_tujuan.add_run("Yth. Kepala Kejaksaan Tinggi Bali\nDi -\n").bold = False
    indent_run = p_tujuan.add_run("\tDenpasar")
    indent_run.bold = False

    # 5. JUDUL SURAT PENGANTAR & NOMOR
    p_judul = doc.add_paragraph()
    p_judul.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_judul.paragraph_format.space_after = Pt(0)
    run_judul = p_judul.add_run("S U R A T  -  P E N G A N T A R")
    run_judul.bold = True
    run_judul.underline = True
    
    p_nomor = doc.add_paragraph()
    p_nomor.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_nomor.paragraph_format.space_after = Pt(10)
    p_nomor.add_run(f"NOMOR : {form_data['nomorTar']}")

    # 6. TABEL UTAMA (Porsi lebar kolom disesuaikan agar NO dan JUMLAH lebih kecil)
    table = doc.add_table(rows=2, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    
    disable_table_autofit(table)
    
    # PERUBAHAN UKURAN: NO (0.8 Cm), NASKAH DINAS (9.2 Cm), JUMLAH (2.0 Cm), KETERANGAN (4.5 Cm)
    col_widths = [Cm(1.5), Cm(8.5), Cm(2.0), Cm(4.5)]
    for i, col in enumerate(table.columns):
        col.width = col_widths[i]

    # Header Tabel
    headers = ["NO", "NASKAH DINAS YANG DIKIRIMKAN", "JUMLAH", "KETERANGAN"]
    for i, title in enumerate(headers):
        cell = table.cell(0, i)
        cell.width = col_widths[i]
        cell.text = title
        set_cell_margins(cell, top=120, bottom=120)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.bold = True
        p.runs[0].font.size = Pt(10)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Isi Baris Tabel
    # Kolom 0: Nomor
    cell_no = table.cell(1, 0)
    cell_no.width = col_widths[0]
    cell_no.text = "1"
    set_cell_margins(cell_no, top=100, bottom=100)
    cell_no.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Kolom 1: Naskah Dinas Yang Dikirimkan
    cell_naskah = table.cell(1, 1)
    cell_naskah.width = col_widths[1]
    set_cell_margins(cell_naskah, top=100, bottom=100, left=150, right=150)
    p_naskah = cell_naskah.paragraphs[0]
    p_naskah.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    r_title = p_naskah.add_run(f"{form_data['nama_laporan']}\n")
    r_title.bold = True
    
    p_naskah.add_run(f"Nomor     :  {form_data['nomorSurat']}\n")
    p_naskah.add_run(f"Tanggal   :  {form_data['tanggal_surat']}\n")
    p_naskah.add_run(f"Perihal     :  {form_data['perihal']}")

    # Kolom 2: Jumlah
    cell_jumlah = table.cell(1, 2)
    cell_jumlah.width = col_widths[2]
    cell_jumlah.text = f"{form_data['jumlah']}"
    set_cell_margins(cell_jumlah, top=100, bottom=100)
    cell_jumlah.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell_jumlah.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Kolom 3: Keterangan
    cell_ket = table.cell(1, 3)
    cell_ket.width = col_widths[3]
    cell_ket.text = f"{form_data['keterangan']}"
    set_cell_margins(cell_ket, top=100, bottom=100, left=100, right=100)
    cell_ket.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    cell_ket.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    p_space = doc.add_paragraph()
    p_space.paragraph_format.space_before = Pt(10)

    # 7. BAGIAN TANDA TANGAN (Sesuai Layout: Baris 1 Kolom 1 = OTENTIKASI)
    ttd_table = doc.add_table(rows=1, cols=2)
    ttd_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    ttd_table.columns[0].width = Cm(7.5)
    ttd_table.columns[1].width = Cm(9.0)

    # Baris 2 Kolom 2: Jabatan & Tanda Tangan (Sisi Kanan)
    cell_pejabat = ttd_table.cell(0, 1)
    p_pejabat = cell_pejabat.paragraphs[0]
    p_pejabat.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_pejabat.add_run(f"{form_data['jabatanTar']}\n\n")
    
    # Input Gambar Esign Digital jika ada
    if os.path.exists("static/images/esign.png"):
        run_esign = p_pejabat.add_run()
        run_esign.add_picture("static/images/esign.png", width=Cm(3.18), height=Cm(1.3))
        p_pejabat.add_run("\n")
    else:
        p_pejabat.add_run("[ E-SIGN PADA DISINI ]\n\n")
    
    run_nama = p_pejabat.add_run(f"{form_data['pejabatTar']}\n")
    run_nama.bold = True
    run_nama.underline = True
    
    p_pejabat.add_run(f"{form_data['nipTar']}")

    # 8. TEMBUSAN
    p_tembusan = doc.add_paragraph()
    p_tembusan.paragraph_format.space_before = Pt(36)
    p_tembusan.paragraph_format.line_spacing = 1.15
    p_tembusan.add_run("Tembusan:\n").underline = True
    
    for idx, item in enumerate(form_data['tembusan'], 1):
        p_tembusan.add_run(f"{idx}.  {item}\n")

    # Simpan Hasil Dokumen Word
    # doc.save(output_path)
    # print(f"Sukses! File berhasil dibuat: {os.path.abspath(output_path)}")
    return doc

MONTHS_ID = [
    "Januari",
    "Februari",
    "Maret",
    "April",
    "Mei",
    "Juni",
    "Juli",
    "Agustus",
    "September",
    "Oktober",
    "November",
    "Desember",
]


def upload_to_drive(doc, file_name, drive_service, target_folder_id=None):
    # Menyimpan file dengan aman di folder /tmp milik Vercel
    local_docx = os.path.join(tempfile.gettempdir(), file_name)
    doc.save(local_docx)

    metadata = {
        "name": file_name,
        "mimeType": "application/vnd.google-apps.document",  # Mengonversi otomatis menjadi Google Docs
    }

    if target_folder_id:
        metadata["parents"] = [target_folder_id]

    media = MediaFileUpload(
        local_docx,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        resumable=True,
    )

    try:
        uploaded = (
            drive_service.files()
            .create(body=metadata, media_body=media, fields="id,webViewLink")
            .execute()
        )
        return uploaded["webViewLink"]
    finally:
        if os.path.exists(local_docx):
            os.remove(local_docx)


def create_pengantar_document(
    drive_service,
    form_data,
    target_folder_id=None,
):
    perihal = form_data["perihal"]
    laporan_date = datetime.datetime.strptime(form_data["tglInput"], "%Y-%m-%d")
    form_data['tembusan'] =  [
            "Yth. Wakil Kepala Kejaksaan Tinggi Bali;",
            "Yth. Asisten Bidang Intelijen Kejaksaan Tinggi Bali;",
            "Yth. Asisten Bidang Pengawasan Kejaksaan Tinggi Bali;",
            "Arsip."
        ]
    form_data['jumlah'] =  "1 (satu)\neks."
    form_data["keterangan"] = "Dengan hormat kami laporkan, selanjutnya mohon petunjuk pimpinan."
    form_data["nama_laporan"] = "LAPORAN INFORMASI KHUSUS"

    tanggal_indonesia = (
        f"{laporan_date.day} "
        f"{MONTHS_ID[laporan_date.month - 1]} "
        f"{laporan_date.year}"
    )
    form_data['tanggal_surat'] = tanggal_indonesia


    file_date = laporan_date.strftime("%Y-%m-%d")

    safe_perihal = sanitize_filename(perihal)

    file_name = f"PENGANTAR-{safe_perihal}"

    dokumen = create_surat_pengantar(form_data)

    try:
        # 4. Panggil fungsi untuk memproses dan mengunggah dokumen
        link_dokumen = upload_to_drive(
            dokumen, file_name, drive_service, target_folder_id=target_folder_id
        )

        # Kembalikan link atau arahkan user ke dokumen Google Docs yang baru jadi
        return link_dokumen

    except Exception as e:
        return f"Terjadi kesalahan saat memproses dokumen: {str(e)}", 500
