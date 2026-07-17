import os
import datetime
import tempfile

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

from googleapiclient.http import MediaFileUpload

from utils.html_utils import html_to_lines, sanitize_filename

MONTHS_ID = [
    "Januari",
    "Februari",
    "Maret",
    "April",
    "Mei",
    "Juni",
    "Juli",
    "Agustus",
    "September",
    "Oktober",
    "November",
    "Desember",
]


def upload_to_drive(doc, file_name, drive_service, target_folder_id=None):
    # Menyimpan file dengan aman di folder /tmp milik Vercel
    local_docx = os.path.join(tempfile.gettempdir(), file_name)
    doc.save(local_docx)

    metadata = {
        "name": file_name,
        "mimeType": "application/vnd.google-apps.document",  # Mengonversi otomatis menjadi Google Docs
    }

    if target_folder_id:
        metadata["parents"] = [target_folder_id]

    media = MediaFileUpload(
        local_docx,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        resumable=True,
    )

    try:
        uploaded = (
            drive_service.files()
            .create(body=metadata, media_body=media, fields="id,webViewLink")
            .execute()
        )
        return uploaded["webViewLink"]
    finally:
        if os.path.exists(local_docx):
            os.remove(local_docx)


def create_lapinhar_document(
    drive_service,
    form_data,
    target_folder_id=None,
):
    nomor_surat = form_data["nomorSurat"]
    perihal = form_data["perihal"]

    kategori = form_data["category"]
    kode_masalah = form_data["kodeMasalah"]

    informasi_html = form_data["infoContent"]
    trend_html = form_data["trendContent"]

    laporan_date = datetime.datetime.strptime(form_data["tglInput"], "%Y-%m-%d")

    pejabat = form_data.get("pejabat", "I Putu Nuriyanto, S.H, M.H.")

    jabatan = form_data.get("jabatan", "Kepala Seksi Intelijen")

    nip = form_data.get("nip", "Jaksa Muda NIP. 19820412 200501 1 003")

    informasi_items = html_to_lines(informasi_html)

    trend_items = html_to_lines(trend_html)

    tanggal_indonesia = (
        f"{laporan_date.day} "
        f"{MONTHS_ID[laporan_date.month - 1]} "
        f"{laporan_date.year}"
    )

    file_date = laporan_date.strftime("%Y-%m-%d")
    mm = laporan_date.strftime("%m")
    yyyy = file_date = laporan_date.strftime("%Y")

    safe_perihal = sanitize_filename(perihal)

    file_name = f"{file_date} - {safe_perihal}"

    local_docx = f"{file_name}.docx"

    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

        header = section.header
        header.paragraphs[0].text = "RAHASIA"
        header.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        footer = section.footer
        footer.paragraphs[0].text = "RAHASIA"
        footer.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("KEJAKSAAN NEGERI TABANAN\n")
    run.bold = True
    run.underline = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run("L.IN.1").bold = True
    p.add_run(f"\n COPY KE: …\nDARI … COPIES")


    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)

    run = p.add_run("LAPORAN INFORMASI HARIAN")
    run.bold = True
    run.underline = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"NOMOR:LIH-{nomor_surat}/N.1.17/{kode_masalah}/{mm}/{yyyy}")


    run.bold = True
    p = doc.add_paragraph()
    p.add_run("I. INFORMASI YANG DIPEROLEH").bold = True
    p.paragraph_format.space_after = Pt(0)

    for item in informasi_items:
        bullet = doc.add_paragraph(style="List Number")
        bullet.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        bullet.add_run(item)

    p = doc.add_paragraph()
    p.add_run("II. SUMBER INFORMASI").bold = True
    p.paragraph_format.space_after = Pt(0)

    doc.add_paragraph(
        "Bahwa informasi yang diperoleh "
        "merupakan organik Tim Intelijen "
        "Kejaksaan Negeri Tabanan."
    )

    p = doc.add_paragraph()
    p.add_run("III. TREND PERKEMBANGAN/PERKIRAAN").bold = True
    p.paragraph_format.space_after = Pt(0)

    for item in trend_items:
        bullet = doc.add_paragraph(style="List Bullet")
        bullet.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        bullet.add_run(item)

    p = doc.add_paragraph()
    p.add_run("IV. SARAN/PENDAPAT").bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(0)

    doc.add_paragraph(
        "Agar Intelijen Kejaksaan Negeri Tabanan tetap berkoordinasi dengan instansi terkait kegiatan yang dilaksanakan di kabupaten Tabanan dan agar dapat dilaporkan kepada pimpinan secara berjenjang."
    ).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    table = doc.add_table(rows=1, cols=2)

    right_cell = table.cell(0, 1)

    p = right_cell.paragraphs[0]

    # rata tengah
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p.add_run(
        f"\nSingasana, {tanggal_indonesia}\n"
        f"{jabatan}\n\n"
    )

    # Tambahkan gambar tanda tangan
    run_gambar = p.add_run()

    run_gambar.add_picture("static/images/esign.png", width=Cm(3.18), height=Cm(1.3))
    nama = p.add_run(f"\n{pejabat}\n")
    nama.bold = True
    nama.underline = True

    p.add_run(f"{nip}")

    try:
        # 4. Panggil fungsi untuk memproses dan mengunggah dokumen
        link_dokumen = upload_to_drive(
            doc, file_name, drive_service, target_folder_id=target_folder_id
        )

        # Kembalikan link atau arahkan user ke dokumen Google Docs yang baru jadi
        return link_dokumen

    except Exception as e:
        return f"Terjadi kesalahan saat memproses dokumen: {str(e)}", 500
